import os
import pytest
from utils.embeddings.embedding_openai_faiss_util import OpenAIFaissUtil
from test_config import TestConfig
from openai import OpenAI
import docx

class TestOpenAIFaissUtil:
    @pytest.fixture
    def faiss_util(self):
        # 使用临时测试文件
        test_index_path = "static/embedding/test_faiss_index"
        # 创建OpenAI客户端并配置
        client = OpenAI(
            api_key=TestConfig.OPENAI_API_KEY,
            base_url=TestConfig.OPENAI_BASE_URL
        )
        util = OpenAIFaissUtil(index_file_path=test_index_path, client=client)
        yield util
        # 测试后清理文件
        # if os.path.exists(test_index_path):
        #     os.remove(test_index_path)
    
    # pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add -v
    def test_add(self,faiss_util):
        # 准备测试数据
        test_texts = [
            "Python是一种流行的编程语言",
            "Java也是一种常用的编程语言",
            "机器学习是人工智能的一个子领域",
            "深度学习是机器学习的一个分支",
            "自然语言处理是人工智能的重要应用",
            "人工智能真的太好玩了",
            "朝辞白帝彩云间，千里江陵一日还。"
        ]
        
        # 测试添加文本
        result = faiss_util.add_texts(test_texts)
        
        # 输出添加结果
        print(f"\n=== 添加文本结果 ===")
        print(f"新添加的文本数量: {result['added_count']}")
        print(f"跳过的重复文本数量: {result['skipped_count']}")
        print(f"数据库中总文本数量: {result['total_count']}")
        print("\n新添加的文本:")
        for i, text in enumerate(result['added_texts'], 1):
            print(f"{i}. {text}")
        

    # pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_search_similar -v
    def test_search_similar(self,faiss_util):
        # 测试相似文本搜索
        query = "赌博"
        results = faiss_util.search_similar(query, top_k=2)
        print("\n搜索结果:")
        print(results)
    
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add_and_search_texts -v
    或
    pytest --capture=no tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add_and_search_texts -v
    """
    def test_add_and_search_texts(self, faiss_util):
        # 准备测试数据
        test_texts = [
            "Python是一种流行的编程语言",
            "Java也是一种常用的编程语言",
            "机器学习是人工智能的一个子领域",
            "深度学习是机器学习的一个分支",
            "自然语言处理是人工智能的重要应用"
        ]
        
        # 测试添加文本
        faiss_util.add_texts(test_texts)
        assert len(faiss_util.texts) == len(test_texts)
        
        # 测试相似文本搜索
        query = "编程语言"
        results = faiss_util.search_similar(query, top_k=2)
        print("\n搜索结果:")
        print(results)
        
        # 验证结果
        assert len(results) == 2
        assert all(isinstance(r['text'], str) for r in results)
        assert all(isinstance(r['score'], float) for r in results)
        assert all(isinstance(r['rank'], int) for r in results)
        
        # 验证排名是否正确
        assert results[0]['rank'] == 1
        assert results[1]['rank'] == 2
        
        # 验证相关性（前两个结果应该是包含"编程语言"的文本）
        programming_related = [r['text'] for r in results]
        assert any("Python" in text for text in programming_related)
        assert any("Java" in text for text in programming_related)
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_search_with_empty_index -v
    """
    def test_search_with_empty_index(self, faiss_util):
        # 测试空索引的搜索行为
        results = faiss_util.search_similar("测试查询", top_k=5)
        print("\n空索引搜索结果:")
        print(results)
        assert len(results) == 0
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add_single_text -v
    """
    def test_add_single_text(self, faiss_util):
        # 测试添加单个文本
        single_text = "这是一个测试文本"
        faiss_util.add_texts([single_text])
        assert len(faiss_util.texts) == 1
        assert faiss_util.texts[0] == single_text
        
        # 验证可以搜索到添加的文本
        results = faiss_util.search_similar("测试", top_k=1)
        print("\n单文本搜索结果:")
        print(results)
        assert len(results) == 1
        assert results[0]['text'] == single_text 
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add_long_text -v
    """
    def test_add_long_text(self, faiss_util):
        # 准备一段长文本
        # long_text = """
        # 人工智能（Artificial Intelligence，简称AI）是计算机科学的一个分支，它企图了解智能的实质，
        # 并生产出一种新的能以人类智能相似的方式做出反应的智能机器。人工智能的研究包括机器人、
        # 语言识别、图像识别、自然语言处理和专家系统等。人工智能从诞生以来，理论和技术日益成熟，
        # 应用领域也不断扩大，可以设想，未来人工智能带来的科技产品，将会是人类智慧的"容器"。
        # 人工智能是对人的意识、思维的信息过程的模拟。人工智能不是人的智能，但能像人那样思考、
        # 也可能超过人的智能。人工智能是一种像人一样思考并采取行动的计算机系统。这种计算机系统
        # 能够从经验中学习，调整到新的输入，并执行类似人类的任务。从简单的分类到视觉感知，从语音
        # 识别到决策制定，人工智能已经并正在改变人类的生活方式。深度学习、机器学习是实现人工智能的一个途径，
        # 人工智能是目标，深度学习是实现目标的手段。
        # """
        long_text="""攀岩高手们咬紧牙关你争我夺，半山腰的时候体能透支全靠意志在硬扛，突然下面有人喊哥们
这边有电梯，这个就是光刻巨头阿斯麦尔的崛起故事
商业是一个漫长的马拉松先发不一定是优势，也可能意味着严重的先行者惩罚
光刻机曾经是尼康的天下，阿斯麦尔是个名不见经传的小厂
别人吃肉他喝一点汤，吃不饱但是也饿不着一直不瘟不火，直到一次契机的出现，我们知道光刻机是生产芯片的，原理就是用光来雕刻电路图，光越细刻的东西就越多芯片就越强
问题是细不下去了，怎么办呢
大家不停地缩短波长，但是到193纳米的时候出现了瓶颈，想缩短到157纳米却发现怎么也无法突破，为了攻克难关几乎整个半导体行业都参与进来，砸了数十亿美金无数的人力物力集在两个方案，一个是尼康的稳健方案，就是还用现在的技术升级到157纳米的F2激光；一个是联盟的激进方案，直接采用全新光源跳过157纳米直接到达10纳米级别，但是商业是要考虑成本的，这两个方案在当时看来，要么成本太高要么难度太大。所有人都要咬牙硬扛的时候，当时的台积电副总经理林本坚，提出了一个天马行空的想法，能不能既不换光源也不升级，不就是把光变细，通过水折射一下不就细了，为什么一定要在空气里传播，193纳米除1.4的折射率等于132纳米远超157纳米
成本低难度低，落地性高设备无需大改，想办法把空气换成水攻克介质难关就行了
更进一步，如果原理跑通的话，换成折射率更高的液体，理论上还能继续辨析，可扩展性也非常好
这下很多人就不开心了，凭什么我们辛辛苦苦爬山，你按个电梯就上来了，凭什么我们投入了海量的人力物力，你改个介质就行了，你不要过来搅局
于是林本健就发表了很多论文来消除疑虑论证可行性，尝试说服各个大厂采用这个以水为介质的方案，但基本上都被拒绝，毕竟大家没日没夜投入了这么多，突然要作废了改走另外一条路，情感上不太能够接受，最后只有阿斯玛尔同意了，一方面他是一个小厂调头方便没有那么多顾虑，巨头们都已经到半山腰了，他还在下面系鞋带，所以就试一下。万一电梯可以用，另一方面是利益驱动，阿斯马尔仔细分析过市场需求，发现搞定新技术的话，拿下两个龙头客户的概率非常大，也就说销量不用愁，于是才决定采用的，结果就是我们看到的阿斯迈尔异军突起
2004年新产品研发成功击败尼康，2009年阿斯迈尔占据70%的份额而尼康则变成了行业小弟
船越大内部阻力就越大，临时调头就越难，自我革命就越不可能
从来没有人规定赛道必须得是一维的，追赶一定得是线性的
看得见的对手是堡垒，看不见的对手是刺客，意识到危险的时候，往往已经就来不及了
 
"""
        
        # 测试添加长文本，设置较小的chunk_tokens以便测试分段
        result = faiss_util.add_long_text(long_text, chunk_tokens=200)
        
        # 输出添加结果
        print(f"\n=== 长文本添加结果 ===")
        print(f"总分段数: {result['total_chunks']}")
        print(f"总token数: {result['total_tokens']}")
        print(f"各段token数: {result['chunks_info']}")
        print(f"新添加的文本数量: {result['added_count']}")
        print(f"跳过的重复文本数量: {result['skipped_count']}")
        print(f"数据库中总文本数量: {result['total_count']}")
        print("\n新添加的文本段:")
        for i, text in enumerate(result['added_texts'], 1):
            print(f"\n段落 {i} (tokens: {result['chunks_info'][i-1]}):")
            print(text)
        
        # 验证结果
        assert result['total_chunks'] > 1  # 确保文本被分段了
        assert result['total_tokens'] > 0  # 确保有token
        assert len(result['chunks_info']) == result['total_chunks']  # 确保每段都有token统计
        assert result['added_count'] == result['total_chunks']  # 确保所有段都被添加
        assert result['skipped_count'] == 0  # 首次添加不应该有重复
        
        # 测试重复添加
        repeat_result = faiss_util.add_long_text(long_text, chunk_tokens=200)
        print(f"\n=== 重复添加结果 ===")
        print(f"跳过的重复文本数量: {repeat_result['skipped_count']}")
        print(f"新添加的文本数量: {repeat_result['added_count']}")
        assert repeat_result['skipped_count'] == repeat_result['total_chunks']  # 所有段都应该被跳过
        assert repeat_result['added_count'] == 0  # 不应该有新添加的文本
        assert repeat_result['added_count'] == 0  # 不应该有新添加的文本 
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_search_fast -v
    """
    def test_search_fast(self, faiss_util):
        # 准备测试数据
        test_texts = [
            "Python是一种流行的编程语言",
            "Java也是一种常用的编程语言",
            "机器学习是人工智能的一个子领域",
            "深度学习是机器学习的一个分支",
            "自然语言处理是人工智能的重要应用"
        ]
        
        # 测试快速搜索
        query = "编程语言"
        results = faiss_util.search_fast(
            query=query, 
            texts=test_texts, 
            top_k=2
        )
        
        print("\n=== 快速搜索结果 ===")
        print(f"查询文本: {query}")
        print(f"返回结果数量: {len(results)}")
        for result in results:
            print(f"\n排名 {result['rank']}:")
            print(f"文本: {result['text']}")
            print(f"相似度得分: {result['score']}")
        
        # 验证结果
        assert len(results) == 2
        assert all(isinstance(r['text'], str) for r in results)
        assert all(isinstance(r['score'], float) for r in results)
        assert all(isinstance(r['rank'], int) for r in results)
        
        # 验证相关性（前两个结果应该是包含"编程语言"的文本）
        programming_related = [r['text'] for r in results]
        assert any("Python" in text for text in programming_related)
        assert any("Java" in text for text in programming_related)
        
        # 测试空列表
        empty_results = faiss_util.search_fast("测试", [])
        assert len(empty_results) == 0 
    
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_embedding_dimension -v
    """
    def test_embedding_dimension(self, faiss_util):
        # 获取向量维度
        dimension = OpenAIFaissUtil.get_embedding_dimension(faiss_util.client)
        
        print(f"\n=== 向量维度测试 ===")
        print(f"模型输出向量维度: {dimension}")
        
        # 获取一个实际的embedding向量并验证其维度
        test_text = "这是一个测试文本"
        embedding = faiss_util.get_embedding(test_text)
        
        print(f"\n=== Embedding详细信息 ===")
        print(f"测试文本: {test_text}")
        print(f"向量维度: {len(embedding)}")
        print(f"向量类型: {type(embedding)}")
        print(f"向量前5个值: {embedding[:5]}")
        print(f"向量后5个值: {embedding[-5:]}")
        
        # 验证维度是否正确（text-embedding-3-small 模型应该是1536维）
        assert dimension == 1536
        assert len(embedding) == dimension
        
        # 验证实例中的维度是否正确
        assert faiss_util.dimension == dimension
        
        # 验证索引的维度是否正确
        assert faiss_util.index.d == dimension 
        
    """运行方法：
    pytest -s tests/utils/embedding_openai_faiss_util_test.py::TestOpenAIFaissUtil::test_add_document -v
    """
    def test_add_document(self, faiss_util):
        # 读取 docx 文件
        doc = docx.Document('tempfiles/1.docx')
        full_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # 使用 add_long_text 方法添加文档内容
        result = faiss_util.add_long_text(full_text, chunk_tokens=500)  # 使用较小的chunk_size便于测试
        
        print("\n=== 文档添加结果 ===")
        print(f"文档被分成了 {result['total_chunks']} 个片段")
        print(f"总共包含 {result['total_tokens']} 个tokens")
        print(f"新添加的文本数量: {result['added_count']}")
        print(f"数据库中总文本数量: {result['total_count']}")
        print("\n各段token数量:")
        for i, token_count in enumerate(result['chunks_info'], 1):
            print(f"第{i}段: {token_count} tokens")
        
        # 验证文档是否被成功添加
        assert result['total_chunks'] > 0  # 确保文档被分段
        assert result['total_tokens'] > 0  # 确保有内容
        assert result['added_count'] == result['total_chunks']  # 确保所有段都被添加
        assert len(result['chunks_info']) == result['total_chunks']  # 确保每段都有token统计
        
        # 测试搜索文档内容
        # 从文档中提取一个短语来搜索
        search_text = doc.paragraphs[0].text[:20]  # 使用第一段的前20个字符
        search_results = faiss_util.search_similar(search_text, top_k=2)
        
        print("\n=== 文档搜索结果 ===")
        print(f"搜索文本: {search_text}")
        for i, result in enumerate(search_results, 1):
            print(f"\n结果 {i}:")
            print(f"文本: {result['text']}")
            print(f"相似度得分: {result['score']}")
        
        # 验证搜索结果
        assert len(search_results) > 0  # 确保能搜索到结果
        

